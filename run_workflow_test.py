import os
import torch
import joblib
import pandas as pd
from transformers import AutoTokenizer, AutoModelForSequenceClassification
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score, confusion_matrix
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np

def extract_head_tail(text, head_words=128, tail_words=384):
    words = str(text).split()
    if len(words) <= 512:
        return str(text)
    return " ".join(words[:head_words] + ["... [TEXT TRUNCATED] ..."] + words[-tail_words:])

def run_evaluation():
    print("Loading Dataset...")
    df = pd.read_csv('legal_bert_v2/legal_text_classification.csv')
    
    model_path = './legal_bert_v2'
    classes = joblib.load(f'{model_path}/label_classes.joblib')
    
    # Sample at least 20 cases (try 2 from each class if available)
    sampled_df = pd.DataFrame()
    for cls in classes:
        # Match case insensitive
        cls_df = df[df['case_outcome'].str.lower() == cls.lower()]
        if not cls_df.empty:
            sampled_df = pd.concat([sampled_df, cls_df.sample(min(3, len(cls_df)), random_state=42)])
    
    # If we don't have enough, just sample randomly to hit 30
    if len(sampled_df) < 30:
        remaining = 30 - len(sampled_df)
        sampled_df = pd.concat([sampled_df, df.sample(remaining, random_state=42)])
        
    sampled_df = sampled_df.sample(frac=1, random_state=42).reset_index(drop=True)
    
    print(f"Sampled {len(sampled_df)} real cases from dataset.")
    
    test_cases = []
    for _, row in sampled_df.iterrows():
        test_cases.append({
            'text': row['case_text'],
            'label': str(row['case_outcome']).lower()
        })

    print("Loading Model...")
    tokenizer = AutoTokenizer.from_pretrained(model_path)
    model = AutoModelForSequenceClassification.from_pretrained(model_path)
    
    device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    model.to(device)
    model.eval()

    y_true = []
    y_pred = []
    
    print("Running Inference on Real Test Cases...")
    for idx, case in enumerate(test_cases):
        text = extract_head_tail(case['text'])
        inputs = tokenizer(text, padding='max_length', truncation=True, max_length=512, return_tensors="pt")
        inputs = {k: v.to(device) for k, v in inputs.items()}
        
        with torch.no_grad():
            outputs = model(**inputs)
            probs = torch.nn.functional.softmax(outputs.logits, dim=1)
            predicted_idx = torch.argmax(probs, dim=1).item()
            prediction = classes[predicted_idx].lower()
            
        y_true.append(case['label'])
        y_pred.append(prediction)
        print(f"Case {idx+1}: True='{case['label']}', Pred='{prediction}'")

    # Metrics Calculation
    acc = accuracy_score(y_true, y_pred)
    precision = precision_score(y_true, y_pred, average='macro', zero_division=0)
    recall = recall_score(y_true, y_pred, average='macro', zero_division=0)
    f1 = f1_score(y_true, y_pred, average='macro', zero_division=0)
    
    print("\n--- METRICS ---")
    print(f"Accuracy: {acc:.4f}")
    print(f"Precision: {precision:.4f}")
    print(f"Recall: {recall:.4f}")
    print(f"F1-Score: {f1:.4f}")

    # Generate Graphs
    print("Generating Graphs...")
    sns.set_theme(style="whitegrid")
    
    metrics_names = ['Accuracy', 'Precision', 'Recall', 'F1-Score']
    metrics_vals = [acc, precision, recall, f1]
    
    plt.figure(figsize=(8, 5))
    ax = sns.barplot(x=metrics_names, y=metrics_vals, hue=metrics_names, palette="viridis", legend=False)
    plt.ylim(0, 1.1)
    plt.title("Model Performance on Real Dataset", fontsize=14)
    for p in ax.patches:
        ax.annotate(format(p.get_height(), '.2f'), 
                    (p.get_x() + p.get_width() / 2., p.get_height()), 
                    ha = 'center', va = 'center', xytext = (0, 9), textcoords = 'offset points')
    plt.tight_layout()
    plt.savefig('metrics_bar_chart.png', dpi=300)
    plt.close()

    cm = confusion_matrix(y_true, y_pred, labels=[c.lower() for c in classes])
    plt.figure(figsize=(10, 8))
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', xticklabels=classes, yticklabels=classes)
    plt.title("Confusion Matrix (Real Data)", fontsize=14)
    plt.ylabel('True Label')
    plt.xlabel('Predicted Label')
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    plt.savefig('confusion_matrix.png', dpi=300)
    plt.close()

    print("Updating Word Document Report...")
    doc = Document()
    doc.add_heading('Legal-BERT v2.0 Updated Evaluation Report', 0)
    doc.add_paragraph('This report contains the evaluation results of the Legal-BERT v2.0 model tested against 30 real cases sampled directly from the training real-world dataset.')
    
    doc.add_heading('1. Overall Performance Metrics', level=1)
    doc.add_paragraph(f"Accuracy: {acc:.2%}")
    doc.add_paragraph(f"Precision (Macro): {precision:.2%}")
    doc.add_paragraph(f"Recall (Macro): {recall:.2%}")
    doc.add_paragraph(f"F1-Score (Macro): {f1:.2%}")
    
    doc.add_picture('metrics_bar_chart.png', width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('2. Confusion Matrix', level=1)
    doc.add_picture('confusion_matrix.png', width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('3. Real Test Cases Sample', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Text Snippet'
    hdr_cells[2].text = 'True Label'
    hdr_cells[3].text = 'Predicted Label'
    
    for idx, case in enumerate(test_cases[:20]): # Show first 20 in table
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[1].text = str(case['text'])[:100] + "..."
        row_cells[2].text = case['label']
        row_cells[3].text = y_pred[idx]

    doc.save('Evaluation_Report.docx')
    print("Report updated as 'Evaluation_Report.docx'")

if __name__ == '__main__':
    run_evaluation()
